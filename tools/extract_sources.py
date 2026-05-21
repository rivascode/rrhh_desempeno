import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


BASE = Path(__file__).resolve().parents[1]
DOCS = BASE / "documentos_base"
OUT = BASE / "analysis"


def clean(value):
    if value is None:
        return ""
    text = str(value).replace("\r", " ").replace("\n", " ")
    return re.sub(r"\s+", " ", text).strip()


def row_values(row):
    return [clean(cell.value) for cell in row]


def compact_row(values):
    return [value for value in values if value]


def analyze_xlsx(path):
    wb = load_workbook(path, data_only=False, read_only=True)
    sheets = []
    for ws in wb.worksheets:
        non_empty = 0
        sample_rows = []
        formula_count = 0
        merged_ranges = [str(rng) for rng in getattr(ws, "merged_cells", []).ranges] if hasattr(ws, "merged_cells") else []
        keyword_hits = []
        for idx, row in enumerate(ws.iter_rows(), start=1):
            values = row_values(row)
            compact = compact_row(values)
            if compact:
                non_empty += 1
                if len(sample_rows) < 25:
                    sample_rows.append({"row": idx, "values": compact[:18]})
                joined = " | ".join(compact)
                if re.search(r"compet|objetiv|desempeñ|clima|pregunta|respuesta|encuesta|factor|indicador|conducta", joined, re.I):
                    keyword_hits.append({"row": idx, "values": compact[:18]})
            for cell in row:
                value = cell.value
                if isinstance(value, str) and value.startswith("="):
                    formula_count += 1
        sheets.append(
            {
                "name": ws.title,
                "max_row": ws.max_row,
                "max_column": ws.max_column,
                "non_empty_rows": non_empty,
                "formula_count": formula_count,
                "merged_ranges_sample": merged_ranges[:20],
                "sample_rows": sample_rows,
                "keyword_hits": keyword_hits[:40],
            }
        )
    return {"file": path.name, "type": "xlsx", "sheets": sheets}


def analyze_docx(path):
    doc = Document(path)
    paragraphs = [clean(p.text) for p in doc.paragraphs]
    paragraphs = [p for p in paragraphs if p]
    tables = []
    for table_idx, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows[:20]:
            rows.append([clean(cell.text) for cell in row.cells])
        tables.append({"index": table_idx, "row_count": len(table.rows), "column_count": len(table.columns), "sample_rows": rows})

    headings = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = clean(p.text)
        if text and ("heading" in style.lower() or "título" in style.lower() or text.isupper()):
            headings.append({"style": style, "text": text})

    keyword_paragraphs = []
    for idx, text in enumerate(paragraphs, start=1):
        if re.search(r"compet|nivel|conducta|evaluaci|desempeñ|diccionario", text, re.I):
            keyword_paragraphs.append({"paragraph": idx, "text": text[:450]})

    frequent_terms = Counter()
    for text in paragraphs:
        for word in re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]{5,}", text.lower()):
            if word not in {"para", "como", "sobre", "entre", "nivel", "debe", "deben", "según"}:
                frequent_terms[word] += 1

    return {
        "file": path.name,
        "type": "docx",
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "headings_sample": headings[:80],
        "paragraphs_sample": paragraphs[:60],
        "keyword_paragraphs": keyword_paragraphs[:120],
        "tables": tables[:20],
        "top_terms": frequent_terms.most_common(40),
    }


def main():
    OUT.mkdir(exist_ok=True)
    results = []
    for path in sorted(DOCS.iterdir()):
        if path.suffix.lower() == ".xlsx":
            results.append(analyze_xlsx(path))
        elif path.suffix.lower() == ".docx":
            results.append(analyze_docx(path))
    output_path = OUT / "source_analysis.json"
    output_path.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    print(output_path)


if __name__ == "__main__":
    sys.exit(main())
